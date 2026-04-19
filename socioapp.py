import streamlit as st
import os
import fitz  # PyMuPDF
import gdown
from docx import Document
from docx.shared import Inches
from io import BytesIO

# --- 1. SOCIOLOGY CONFIGURATION (BRUNEI ZONE 4) ---
# The app will now look into that 'Secrets' box you just filled!
# if password == st.secrets["ADMIN_PASSWORD"]:
#    st.success("Admin Access Granted")
# ... rest of your admin code

SYLLABUS_CODE = "9699"
#GD_FOLDER_ID = "YOUR_SOCIOLOGY_GD_FOLDER_ID_HERE"
GD_FOLDER_ID = "169rNmys_Uh_fL1ZY1cxlPCdlhBiXxlj8"

FOLDERS = {
    "June QP": "9699_June_qp",
    "Nov QP": "9699_Nov_qp",
    "June MS": "9699_June_ms",
    "Nov MS": "9699_Nov_ms"
}

for folder in FOLDERS.values():
    if not os.path.exists(folder):
        os.makedirs(folder)


def sync_from_drive():
    try:
        with st.spinner("🔄 Syncing Sociology Library..."):
            gdown.download_folder(id=GD_FOLDER_ID, output=".", quiet=True)
        st.success("✅ Sociology Library Updated!")
    except Exception as e:
        st.error(f"Sync Error: {e}")


# --- 2. APP STATE ---
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []
if 'search_results' not in st.session_state:
    st.session_state.search_results = []

# --- 3. UI LAYOUT ---
st.set_page_config(page_title="9699 Sociology Handout Builder", layout="wide")
st.title(" PUSAT TINGKATAN ENAM SENGKURONG ")
st.subheader("📚 9699 Sociology PYP Handout Platform (Brunei Variant)")

with st.sidebar:
    st.header("Cloud Controls")
    if st.button("🔄 Sync New Files"):
        sync_from_drive()
    st.info("Upload Sociology PDFs to Google Drive, then sync here.")


def get_filename_pattern(month, year, paper_type, paper_num):
    short_year = year[-2:]
    # June is 's' (Summer), Nov is 'w' (Winter)
    month_code = 's' if month == "June" else 'w'
    return f"9699_{month_code}{short_year}_{paper_type}_{paper_num}"


def search_pdfs(keyword_list, folder_path):
    results = []
    if not os.path.exists(folder_path): return results
    for file in os.listdir(folder_path):
        if file.endswith(".pdf"):
            try:
                doc = fitz.open(os.path.join(folder_path, file))
                for page_num in range(len(doc)):
                    text = doc[page_num].get_text().lower()
                    if all(k.lower() in text for k in keyword_list):
                        results.append({
                            "file": file, "page": page_num, "path": os.path.join(folder_path, file)
                        })
                doc.close()
            except:
                continue
    return results


tab1, tab2, tab3, tab4 = st.tabs(["🔍 Search & Extract", "📅 View Exam Papers", "📝 Export Handout", "⚙️ Admin"])

# --- TAB 1: SEARCH ---
with tab1:
    st.header("Search Sociology Topics")
    keywords = st.text_input("Enter keywords (e.g., 'Marxism', 'Functionalism', 'Domestic division of labour')")

    if st.button("Search Papers", type="primary"):
        if keywords:
            with st.spinner("Scanning sociology papers..."):
                all_results = []
                for folder_path in FOLDERS.values():
                    all_results += search_pdfs([keywords], folder_path)
                st.session_state.search_results = all_results
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.search_results:
        st.write(f"Found {len(st.session_state.search_results)} matching pages:")
        for idx, item in enumerate(st.session_state.search_results):
            c1, c2 = st.columns([4, 1])
            c1.write(f"📄 **{item['file']}** (Page {item['page'] + 1})")
            if c2.button("➕ Add", key=f"add_{idx}"):
                st.session_state.handout_basket.append(item)
                st.toast("Added to basket!")

# --- TAB 2: VIEW FULL PAPERS (BRUNEI SPECIFIC) ---
with tab2:
    st.header("Download Full Papers (Zone 4)")
    c1, c2, c3 = st.columns(3)
    with c1:
        v_year = st.selectbox("Year", [str(y) for y in range(2026, 2018, -1)])
    with c2:
        v_month = st.selectbox("Month", ["June", "Nov"])
    with c3:
        # Prioritizing Variant 3 for JUNE and variant 2 for November)
        v_paper = st.selectbox("Paper Component", ["12", "22", "32", "42", "13", "23", "33", "43"])

    qp_name = get_filename_pattern(v_month, v_year, "qp", v_paper) + ".pdf"
    ms_name = get_filename_pattern(v_month, v_year, "ms", v_paper) + ".pdf"

    col_q, col_m = st.columns(2)
    with col_q:
        path = os.path.join(FOLDERS[f"{v_month} QP"], qp_name)
        if os.path.exists(path):
            st.success(f"Found QP: {qp_name}")
            with open(path, "rb") as f:
                st.download_button("Download QP", f, file_name=qp_name)
        else:
            st.info(f"Note: {qp_name} not in local library.")

    with col_m:
        path_ms = os.path.join(FOLDERS[f"{v_month} MS"], ms_name)
        if os.path.exists(path_ms):
            st.success(f"Found MS: {ms_name}")
            with open(path_ms, "rb") as f:
                st.download_button("Download MS", f, file_name=ms_name)

# --- TAB 3: EXPORT ---
with tab3:
    st.header("Generate Sociology Worksheet")
    if st.session_state.handout_basket:
        st.write(f"Selected Items: **{len(st.session_state.handout_basket)}**")
        if st.button("🗑️ Clear Basket"):
            st.session_state.handout_basket = []
            st.rerun()

        if st.button("🪄 Generate Word Document", type="primary"):
            doc = Document()
            doc.add_heading('PTES 9699 Sociology Handout', 0)
            for item in st.session_state.handout_basket:
                doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
                pdf_doc = fitz.open(item['path'])
                page = pdf_doc.load_page(item['page'])
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_data = BytesIO(pix.tobytes("png"))
                doc.add_picture(img_data, width=Inches(6))
                doc.add_page_break()
                pdf_doc.close()

            target = "9699_Sociology_Handout.docx"
            doc.save(target)
            with open(target, "rb") as f:
                st.download_button("📥 Click to Download (.docx)", f, file_name=target)
    else:
        st.info("Your basket is empty. Use the Search tab to add questions.")

# --- TAB 4: ADMIN ---
with tab4:
    st.header("Admin Panel")
    pwd = st.text_input("Password", type="password")
    if pwd == st.secrets["ADMIN_PASSWORD"]:
        st.success("Admin Access Granted")
        u_col, d_col = st.columns(2)
        with u_col:
            st.subheader("Manual Upload")
            dest = st.selectbox("Folder", list(FOLDERS.keys()))
            up_files = st.file_uploader("Select PDFs", type="pdf", accept_multiple_files=True)
            if st.button("Upload"):
                for f in up_files:
                    with open(os.path.join(FOLDERS[dest], f.name), "wb") as s: s.write(f.getbuffer())
                st.success("Files saved!")
        with d_col:
            st.subheader("Management")
            dest_d = st.selectbox("Target Folder", list(FOLDERS.keys()))
            files_in_folder = os.listdir(FOLDERS[dest_d])
            to_del = st.selectbox("Delete File", ["---"] + files_in_folder)
            if to_del != "---" and st.button("Confirm Delete"):
                os.remove(os.path.join(FOLDERS[dest_d], to_del))
                st.rerun()

# --- FOOTER ---
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; width: 100%;">
        <p style="font-size: 20px; font-weight: bold; margin-bottom: 5px;">
            ✨ PTES 9699 Sociology Resource Portal ✨
        </p>
        <p style="font-size: 16px; font-weight: bold; letter-spacing: 0.5px;">
            <span style="color: #FF0000;">🔴 Academically Excellence</span> | 
            <span style="color: #FFD700;">🟡 Future Readiness</span> | 
            <span style="color: #0070FF;">🔵 Digital & Integrity</span> | 
            <span style="color: #28A745;">🟢 Holistic & Growth</span>
        </p>
        <p style="color: gray; font-size: 14px; margin-top: 10px;">
            Creator: Miss Hajah Nurul Haziqah HN (Computer Science PTES)
        </p>
    </div>
    """,
    unsafe_allow_html=True
)
